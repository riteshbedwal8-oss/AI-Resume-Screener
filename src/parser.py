# src/parser.py
# Purpose: Extract raw text from PDF and DOCX resume files

import pdfplumber          # For reading PDF files
import docx                # For reading Word files (python-docx)
import os

def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract all text from a PDF file page by page.
    Returns a single string with all content.
    """
    text = ""
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:                    # Some pages may be images (scanned)
                    text += page_text + "\n"
    except Exception as e:
        print(f"[ERROR] Could not read PDF {file_path}: {e}")
    return text.strip()


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract all paragraphs from a DOCX file.
    Tables are also extracted for completeness.
    """
    text = ""
    try:
        doc = docx.Document(file_path)
        # Extract paragraphs
        for para in doc.paragraphs:
            text += para.text + "\n"
        # Extract text from tables (many resumes use table layouts)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
            text += "\n"
    except Exception as e:
        print(f"[ERROR] Could not read DOCX {file_path}: {e}")
    return text.strip()


def parse_resume(file_path: str) -> str:
    """
    Main function: detects file type and routes to correct parser.
    Accepts .pdf or .docx files.
    """
    ext = os.path.splitext(file_path)[1].lower()   # Get file extension

    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext in [".docx", ".doc"]:
        return extract_text_from_docx(file_path)
    else:
        print(f"[WARNING] Unsupported file format: {ext}")
        return ""


# --- Quick test (run this file directly to verify) ---
if __name__ == "__main__":
    # Replace with an actual resume path to test
    test_path = "data/resumes/sample_resume.pdf"
    if os.path.exists(test_path):
        result = parse_resume(test_path)
        print(result[:500])   # Print first 500 chars
    else:
        print("Place a sample resume in data/resumes/ to test.")